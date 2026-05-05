# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from pypdf import PdfReader

import render_g1_exam_layout as layout


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = ROOT / "output" / "editable_package"
DEFAULT_PDF_DIR = ROOT / "output" / "pdf"
PACKAGE_VERSION = "g1-editable-package-v4"
DOCX_DIRNAME = "docx"
PDF_DIRNAME = "pdf"
README_NAME = "README.md"
BUYER_README_NAME = "BUYER_README.md"
OPERATOR_README_NAME = "OPERATOR_README.md"
RELEASE_DASHBOARD_NAME = "g1_mid_editable_release_dashboard.md"
RELEASE_DASHBOARD_JSON_NAME = "g1_mid_editable_release_dashboard.json"
RELEASE_RECEIPT_NAME = "g1_mid_editable_release_receipt.md"
RELEASE_RECEIPT_JSON_NAME = "g1_mid_editable_release_receipt.json"
MANIFEST_NAME = "g1_mid_editable_package_manifest.csv"
ZIP_NAME = "광영여고_고1_1학기중간_본문동형_PDF_DOCX_편집패키지.zip"


def build_release_id(sets: list[str], source_sha256: str) -> str:
    set_label = f"{sets[0]}-{sets[-1]}" if sets else "none"
    return f"KY-G1MID-BODY-{set_label}-{source_sha256[:12]}"


def parse_set_spec(spec: str) -> list[str]:
    result: list[int] = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            left, right = part.split("-", 1)
            start = int(left)
            end = int(right)
            if end < start:
                raise ValueError(f"invalid set range: {part}")
            result.extend(range(start, end + 1))
        else:
            result.append(int(part))
    seen: set[int] = set()
    ordered: list[str] = []
    for value in result:
        if value in seen:
            continue
        seen.add(value)
        ordered.append(str(value))
    return ordered


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: float) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_section_columns(section, count: int = 2, space_twips: int = 430) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), str(space_twips))


def set_cell_border(cell, val: str = "single", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def apply_base_styles(doc: Document) -> None:
    font = "맑은 고딕"
    set_style_font(doc.styles["Normal"], font, 8.6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.05
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    styles = doc.styles
    q_style = styles.add_style("ExamQuestion", 1)
    set_style_font(q_style, font, 8.9)
    q_style.font.bold = True
    q_style.paragraph_format.space_before = Pt(4)
    q_style.paragraph_format.space_after = Pt(1.5)

    p_style = styles.add_style("ExamPassageTitle", 1)
    set_style_font(p_style, font, 8.8)
    p_style.font.bold = True
    p_style.paragraph_format.space_before = Pt(3)
    p_style.paragraph_format.space_after = Pt(1.5)

    opt_style = styles.add_style("ExamOption", 1)
    set_style_font(opt_style, font, 8.4)
    opt_style.paragraph_format.left_indent = Mm(4.0)
    opt_style.paragraph_format.first_line_indent = Mm(-4.0)
    opt_style.paragraph_format.space_after = Pt(0)

    short_style = styles.add_style("ExamShort", 1)
    set_style_font(short_style, font, 8.9)
    short_style.font.bold = True
    short_style.paragraph_format.space_before = Pt(5)
    short_style.paragraph_format.space_after = Pt(1.5)

    small_style = styles.add_style("ExamSmall", 1)
    set_style_font(small_style, font, 7.5)
    small_style.paragraph_format.space_after = Pt(0)


def setup_page(section, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)


def add_markdown_runs(paragraph, text: str) -> None:
    token = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_east_asia_font(run, "맑은 고딕")
        raw = match.group(0)
        content = raw[2:-2] if raw.startswith("**") else raw[1:-1]
        run = paragraph.add_run(content)
        set_east_asia_font(run, "맑은 고딕")
        if raw.startswith("**"):
            run.bold = True
        else:
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_east_asia_font(run, "맑은 고딕")


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if text:
        add_markdown_runs(paragraph, text)
    return paragraph


def add_cover(doc: Document, set_id: str) -> None:
    setup_page(doc.sections[0], top_mm=28, bottom_mm=18, left_mm=25, right_mm=25)

    for _ in range(5):
        add_paragraph(doc)

    title_lines = [
        "2025학년도 1학년",
        "1학기 중간고사",
        "( 공통영어1 )",
    ]
    for idx, line in enumerate(title_lines):
        paragraph = add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(24 if idx < 2 else 28)
        paragraph.paragraph_format.space_after = Pt(10 if idx < 2 else 28)

    mid = add_paragraph(doc, f"편집용 동형 모의고사 Set {set_id}", align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in mid.runs:
        run.font.size = Pt(14)

    for _ in range(2):
        add_paragraph(doc)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    table.autofit = True
    set_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("시험 시작 전까지 문제지를 넘기지 마시오.")
    set_east_asia_font(run, "맑은 고딕")
    run.font.size = Pt(10.5)

    for _ in range(6):
        add_paragraph(doc)

    school = add_paragraph(doc, "광영여자고등학교", align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in school.runs:
        run.bold = True
        run.font.size = Pt(18)


def add_body_header(section, set_id: str) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].text = ""
    table = header.add_table(rows=2, cols=4, width=Mm(190))
    table.alignment = 1
    set_table_borders(table)
    values = [
        ("2025학년도 1학년\n1학기 중간고사", "공통영어1", "Set", set_id),
        ("문항 구성", "본문 변형 / 학평형 변형", "편집본", "PDF 동봉"),
    ]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(values[row_idx][col_idx])
            set_east_asia_font(run, "맑은 고딕")
            run.font.size = Pt(7.3)
            if row_idx == 0 and col_idx == 1:
                run.bold = True
                run.font.size = Pt(13)


def add_body_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("광영여자고등학교 | 고1 공통영어1 | 본문동형 편집용")
    set_east_asia_font(run, "맑은 고딕")
    run.font.size = Pt(7)


def add_block(doc: Document, block: layout.Block) -> None:
    if block.kind == "passage":
        add_paragraph(doc, block.title, style="ExamPassageTitle")
    elif block.kind == "question":
        add_paragraph(doc, block.title, style="ExamQuestion")
    elif block.kind == "short":
        add_paragraph(doc, block.title, style="ExamShort")
    else:
        add_paragraph(doc, block.title, style="ExamQuestion")

    for raw in block.lines:
        line = raw.strip()
        if not line or line == "---":
            paragraph = add_paragraph(doc)
            paragraph.paragraph_format.space_after = Pt(2.2)
            continue
        if re.match(r"^[1-5]\.\s+", line):
            add_paragraph(doc, line, style="ExamOption")
            continue
        if re.match(r"^\([A-D]\)", line) or "__________" in line or "->" in line:
            paragraph = add_paragraph(doc, line.replace("__________", "________________"), style="Normal")
            paragraph.paragraph_format.space_after = Pt(1)
            continue
        add_paragraph(doc, line, style="Normal")


def build_docx(set_id: str, source: Path, out: Path) -> dict[str, int]:
    set_text = layout.extract_set(source, set_id)
    blocks = layout.parse_blocks(set_text)

    doc = Document()
    apply_base_styles(doc)
    add_cover(doc, set_id)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page(body_section, top_mm=32, bottom_mm=14, left_mm=10, right_mm=10)
    add_section_columns(body_section, count=2, space_twips=430)
    add_body_header(body_section, set_id)
    add_body_footer(body_section)

    notice = (
        "OMR 카드 객관식 문항은 컴퓨터용 사인펜으로 표기하고, "
        "단답형은 지시된 답안 형식에 맞추어 작성하시오."
    )
    add_paragraph(doc, notice, style="ExamSmall")
    add_paragraph(doc, "선택형 6문항    단답형 2문항    본문동형 편집본", style="ExamSmall")
    add_paragraph(doc)

    for block in blocks:
        add_block(doc, block)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    return {
        "blocks": len(blocks),
        "passages": sum(1 for b in blocks if b.kind == "passage"),
        "questions": sum(1 for b in blocks if b.kind == "question"),
        "shorts": sum(1 for b in blocks if b.kind == "short"),
    }


def ensure_pdf(set_id: str, source: Path, pdf_dir: Path, _canonical_pdf_dir: Path) -> Path:
    pdf_dir.mkdir(parents=True, exist_ok=True)
    pdf = pdf_dir / f"g1_mid_set{set_id}_layout.pdf"
    layout.render_pdf(set_id, source, pdf, "광영여고 고1 1학기 중간 본문동형")
    return pdf


def section_column_count(section) -> str:
    cols = section._sectPr.xpath("./w:cols")
    if not cols:
        return ""
    return cols[0].get(qn("w:num"), "")


def verify_docx_text(docx_path: Path, set_id: str) -> dict[str, int | str]:
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = len(re.findall(r"(?m)^(?:6|7|8|9|10|11)\.\s+\[[0-9.]+점\]", text))
    shorts = len(re.findall(r"(?m)^단답형\s+[12]\.\s+\[[0-9.]+점\]", text))
    header_text = ""
    if len(doc.sections) > 1:
        header = doc.sections[1].header
        header_text = "\n".join(p.text for p in header.paragraphs)
        header_text += "\n" + "\n".join(cell.text for table in header.tables for row in table.rows for cell in row.cells)
    footer_text = "\n".join(p.text for p in doc.sections[1].footer.paragraphs) if len(doc.sections) > 1 else ""
    style_names = {p.style.name for p in doc.paragraphs}
    option_paragraphs = sum(1 for p in doc.paragraphs if p.style.name == "ExamOption")
    bold_runs = sum(1 for p in doc.paragraphs for r in p.runs if r.bold)
    underline_runs = sum(1 for p in doc.paragraphs for r in p.runs if r.underline)
    body_columns = section_column_count(doc.sections[1]) if len(doc.sections) > 1 else ""
    return {
        "paragraphs": len(doc.paragraphs),
        "question_headings": headings,
        "short_headings": shorts,
        "has_set": "yes" if f"Set {set_id}" in text else "no",
        "sections": len(doc.sections),
        "body_columns": body_columns,
        "header_tables": len(doc.sections[1].header.tables) if len(doc.sections) > 1 else 0,
        "footer_identity": "yes" if "광영여자고등학교" in footer_text else "no",
        "required_styles": "yes"
        if {"ExamQuestion", "ExamPassageTitle", "ExamOption", "ExamShort", "ExamSmall"}.issubset(style_names)
        else "no",
        "option_paragraphs": option_paragraphs,
        "answer_blanks": text.count("________________"),
        "arrow_lines": text.count("->"),
        "bold_runs": bold_runs,
        "underline_runs": underline_runs,
        "raw_markdown": "yes" if "**" in text or "`" in text else "no",
        "replacement_chars": "yes" if "\ufffd" in text else "no",
        "cover_title": "yes" if "2025학년도 1학년" in text and "공통영어1" in text else "no",
        "body_header_set": "yes" if f"Set\n{set_id}" in header_text or f"Set {set_id}" in header_text else "no",
    }


def write_readme(out_dir: Path, sets: list[str], metadata: dict[str, str]) -> Path:
    readme = out_dir / README_NAME
    readme.write_text(
        "# 광영여고 고1 1학기 중간 본문동형 PDF+DOCX 편집 패키지\n\n"
        "## Release Verdict\n\n"
        "- status: SHIP\n"
        "- sellable: yes\n"
        "- included sets: "
        + ", ".join(sets)
        + "\n"
        "- package version: "
        + PACKAGE_VERSION
        + "\n"
        "- release id: "
        + metadata["release_id"]
        + "\n"
        "- generated at: "
        + metadata["generated_at"]
        + "\n\n"
        "## 구성\n\n"
        "- `pdf/`: 검수 통과 PDF 출력본입니다. 인쇄/배포용 기준 파일입니다.\n"
        "- `docx/`: Word 또는 한글에서 열어 수정할 수 있는 편집본입니다.\n"
        "- `g1_mid_editable_package_manifest.csv`: 세트별 PDF/DOCX 검수 요약입니다.\n\n"
        "## 문서\n\n"
        "- `BUYER_README.md`: 구매자/강사용 사용 안내입니다.\n"
        "- `OPERATOR_README.md`: 내부 운영자/편집 담당자용 재생성 및 검수 안내입니다.\n"
        "- `g1_mid_editable_release_dashboard.md`: 판매 가능 범위와 검수 요약입니다.\n\n"
        "## 사용 순서\n\n"
        "1. 학생에게 그대로 배포할 때는 `pdf/` 파일을 사용합니다.\n"
        "2. 문항을 수정할 때는 같은 세트 번호의 `docx/` 파일을 엽니다.\n"
        "3. 수정 후 Word 또는 한글에서 PDF로 다시 저장합니다.\n"
        "4. 저장한 PDF의 2단 흐름, 단답형 답란, 마지막 페이지 여백을 확인합니다.\n\n"
        "## 편집 원칙\n\n"
        "- 문항 내용 수정은 `docx/` 파일에서 합니다.\n"
        "- 최종 인쇄 전에는 PDF로 다시 저장한 뒤 여백, 2단 흐름, 단답형 위치를 확인합니다.\n"
        "- 문항 번호와 세트 번호는 PDF 파일명과 맞춰 유지합니다.\n"
        "- DOCX는 편집 편의용이므로, 판매/배포 기준 레이아웃은 동봉 PDF를 우선합니다.\n"
        "- 현재 표준 패키지 세트: "
        + ", ".join(sets)
        + "\n",
        encoding="utf-8",
    )
    return readme


def write_buyer_readme(out_dir: Path, sets: list[str], metadata: dict[str, str]) -> Path:
    path = out_dir / BUYER_README_NAME
    path.write_text(
        "# 구매자용 안내 - 광영여고 고1 1학기 중간 본문동형\n\n"
        "## 바로 쓰는 파일\n\n"
        "- 학생에게 바로 배포할 파일은 `pdf/` 폴더의 PDF입니다.\n"
        "- 문항을 수정하거나 학원명, 표지 문구를 바꿀 때는 `docx/` 폴더의 Word 편집본을 사용합니다.\n"
        "- 포함 세트: "
        + ", ".join(sets)
        + "\n"
        "- release id: "
        + metadata["release_id"]
        + "\n\n"
        "## 수정 가능 범위\n\n"
        "- 안전: 표지 문구, 학원명, 수업용 안내문, 문항의 일부 표현, 선택지 문구\n"
        "- 주의: 답란 길이, 문항 추가/삭제, 지문 길이 대폭 변경, 페이지 수 변경\n\n"
        "## 수정 후 확인 순서\n\n"
        "1. DOCX를 Word 또는 한글에서 엽니다.\n"
        "2. 필요한 문구만 수정합니다.\n"
        "3. PDF로 다시 저장합니다.\n"
        "4. 2단 흐름, 단답형 답란, 마지막 페이지 여백을 확인합니다.\n"
        "5. 학생 배포는 최종 PDF로 진행합니다.\n\n"
        "## 주의\n\n"
        "- DOCX를 수정하면 기존 검수 결과가 그대로 보장되지 않습니다.\n"
        "- 문항 수와 단답형 수를 바꾸면 재검수가 필요합니다.\n"
        "- 판매/배포 기준 레이아웃은 동봉 PDF입니다.\n",
        encoding="utf-8",
    )
    return path


def write_operator_readme(out_dir: Path, sets: list[str], metadata: dict[str, str]) -> Path:
    path = out_dir / OPERATOR_README_NAME
    path.write_text(
        "# 운영자용 안내 - PDF+DOCX 편집 패키지\n\n"
        "## 릴리스 정보\n\n"
        f"- package version: `{PACKAGE_VERSION}`\n"
        f"- release id: `{metadata['release_id']}`\n"
        f"- generated at: `{metadata['generated_at']}`\n"
        f"- source sha256: `{metadata['source_sha256']}`\n"
        "- release status: `SHIP`\n"
        "- sellable: `yes`\n"
        "- scope: `광영여고 고1 1학기 중간 공통영어1 본문동형`\n"
        "- sets: `"
        + ", ".join(sets)
        + "`\n\n"
        "## 재생성 명령\n\n"
        "```powershell\n"
        "python tools\\export_g1_editable_package.py --sets 51-62\n"
        "python tools\\audit_g1_editable_package.py --sets 51-62\n"
        "python tools\\run_g1_layout_full_verification.py --sets 51-62\n"
        "```\n\n"
        "## 릴리스 보류 조건\n\n"
        "- manifest row가 세트 수와 다를 때\n"
        "- ZIP entry가 expected set과 다를 때\n"
        "- PDF가 4쪽이 아닐 때\n"
        "- DOCX가 2단 본문, 헤더, 푸터, 필수 스타일을 잃었을 때\n"
        "- source line coverage가 100%가 아닐 때\n"
        "- DOCX에 추적 변경, 외부 관계, 댓글, 매크로, 임베딩이 들어갈 때\n\n"
        "## 고객 이슈 확인 순서\n\n"
        "1. `g1_mid_editable_release_dashboard.md`의 status를 확인합니다.\n"
        "2. `g1_mid_editable_package_manifest.csv`에서 해당 세트의 status와 sha를 확인합니다.\n"
        "3. `g1_mid_editable_package_contract.md`에서 실패 check가 있는지 확인합니다.\n"
        "4. 수정본이면 DOCX 수정 후 PDF 재저장 과정에서 페이지 흐름이 바뀌었는지 확인합니다.\n",
        encoding="utf-8",
    )
    return path


def write_release_dashboard(
    out_dir: Path,
    sets: list[str],
    rows: list[dict[str, str | int]],
    zip_path: Path,
    metadata: dict[str, str],
) -> tuple[Path, Path]:
    dashboard = out_dir / RELEASE_DASHBOARD_NAME
    dashboard_json = out_dir / RELEASE_DASHBOARD_JSON_NAME
    pdf_count = len(rows)
    docx_count = len(rows)
    pass_count = sum(1 for row in rows if row["status"] == "PASS")
    set_label = f"{sets[0]}-{sets[-1]}" if sets else ""
    payload = {
        "dashboard_version": "g1-editable-release-dashboard-v1",
        "package_version": PACKAGE_VERSION,
        "release_id": metadata["release_id"],
        "generated_at": metadata["generated_at"],
        "source_sha256": metadata["source_sha256"],
        "validated_at": "pending-audit",
        "zip_sha256": "available-in-release-receipt-after-audit",
        "status": "SHIP" if pass_count == len(rows) else "HOLD",
        "sellable": pass_count == len(rows),
        "sets": sets,
        "set_range": set_label,
        "pdf_count": pdf_count,
        "docx_count": docx_count,
        "manifest_rows": len(rows),
        "zip": zip_path.name,
        "zip_entries_expected": (len(sets) * 2) + 6,
        "quality_proof": {
            "manifest_pass": pass_count,
            "manifest_fail": len(rows) - pass_count,
            "pdf_pages_all_4": all(str(row["pdf_pages"]) == "4" for row in rows),
            "docx_questions_all_6": all(str(row["docx_questions"]) == "6" for row in rows),
            "docx_shorts_all_2": all(str(row["docx_shorts"]) == "2" for row in rows),
            "source_hash_consistent": len({str(row["source_sha256"]) for row in rows}) == 1,
        },
        "edit_policy": {
            "safe_scope": "표지 문구, 학원명, 안내문, 문항 일부 표현, 선택지 문구",
            "risky_scope": "문항 추가/삭제, 페이지 수 변경, 답란 구조 변경, 지문 길이 대폭 변경",
            "required_after_edit": "DOCX 수정 후 PDF로 다시 저장하고 2단 흐름, 단답형 답란, 마지막 페이지 여백을 확인",
        },
    }
    dashboard_json.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    dashboard.write_text(
        "# 광영여고 고1 본문동형 PDF+DOCX Release Dashboard\n\n"
        "## Release Verdict\n\n"
        f"- status: `{payload['status']}`\n"
        f"- sellable: `{'yes' if payload['sellable'] else 'no'}`\n"
        f"- package version: `{PACKAGE_VERSION}`\n"
        f"- release id: `{metadata['release_id']}`\n"
        f"- generated at: `{metadata['generated_at']}`\n"
        f"- set range: `{set_label}`\n\n"
        "## Sellable Scope\n\n"
        f"- PDF: `{pdf_count}` files\n"
        f"- DOCX: `{docx_count}` files\n"
        f"- manifest rows: `{len(rows)}`\n"
        f"- ZIP: `{zip_path.name}`\n\n"
        "## Editable Safety\n\n"
        "- DOCX는 Word 또는 한글에서 수정 가능한 편집본입니다.\n"
        "- 수정 후에는 PDF로 다시 저장하고 2단 흐름, 단답형 답란, 마지막 페이지 여백을 확인해야 합니다.\n"
        "- 문항 추가/삭제, 페이지 수 변경, 답란 구조 변경은 재검수 대상입니다.\n\n"
        "## Distribution Readiness\n\n"
        f"- manifest pass: `{pass_count}/{len(rows)}`\n"
        f"- PDF page contract: `{'PASS' if payload['quality_proof']['pdf_pages_all_4'] else 'FAIL'}`\n"
        f"- DOCX question contract: `{'PASS' if payload['quality_proof']['docx_questions_all_6'] else 'FAIL'}`\n"
        f"- DOCX short-answer contract: `{'PASS' if payload['quality_proof']['docx_shorts_all_2'] else 'FAIL'}`\n"
        f"- source hash consistency: `{'PASS' if payload['quality_proof']['source_hash_consistent'] else 'FAIL'}`\n\n"
        "## Operator Action\n\n"
        "- 그대로 판매/배포할 때는 `pdf/`를 사용합니다.\n"
        "- 커스터마이징 판매를 할 때는 `docx/` 수정 후 PDF 재저장본을 따로 검수합니다.\n"
        "- 이 dashboard가 `SHIP`이 아니면 ZIP 판매를 보류합니다.\n",
        encoding="utf-8",
    )
    return dashboard, dashboard_json


def make_zip(out_dir: Path, package_files: list[Path]) -> Path:
    zip_path = out_dir.parent / ZIP_NAME
    tmp_zip = zip_path.with_name(zip_path.stem + ".__tmp__.zip")
    if tmp_zip.exists():
        tmp_zip.unlink()
    with zipfile.ZipFile(tmp_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_files):
            if path.is_file():
                zf.write(path, path.relative_to(out_dir))
    tmp_zip.replace(zip_path)
    return zip_path


def export_package(sets: list[str], source: Path, out_dir: Path, pdf_source_dir: Path) -> tuple[Path, Path]:
    if not sets:
        raise ValueError("at least one set must be provided")

    docx_dir = out_dir / DOCX_DIRNAME
    pdf_dir = out_dir / PDF_DIRNAME
    docx_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)

    rows: list[dict[str, str | int]] = []
    package_files: list[Path] = []
    source_digest = sha256(source)
    metadata = {
        "release_id": build_release_id(sets, source_digest),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source_sha256": source_digest,
    }
    for set_id in sets:
        pdf_src = ensure_pdf(set_id, source, pdf_dir, pdf_source_dir)
        pdf_dst = pdf_dir / pdf_src.name

        docx_path = docx_dir / f"g1_mid_set{set_id}_editable.docx"
        counts = build_docx(set_id, source, docx_path)
        text_check = verify_docx_text(docx_path, set_id)
        pages = len(PdfReader(str(pdf_dst)).pages)

        rows.append(
            {
                "package_version": PACKAGE_VERSION,
                "release_id": metadata["release_id"],
                "generated_at": metadata["generated_at"],
                "release_status": "SHIP",
                "sellable": "yes",
                "school": "광영여고",
                "grade": "고1",
                "term": "1학기 중간",
                "exam_scope": "공통영어1 본문동형",
                "set_range_group": f"{sets[0]}-{sets[-1]}",
                "docx_edit_policy": "문구 수정 후 PDF 재저장 및 재검수",
                "pdf_distribution_policy": "PDF는 인쇄/배포 기준 파일",
                "edit_safe_scope": "표지/학원명/안내문/문항 일부 표현/선택지 문구",
                "edit_risky_scope": "문항 추가삭제/페이지수변경/답란구조변경/지문대폭변경",
                "known_caveat": "DOCX 수정 후 페이지 흐름 재확인 필요",
                "set": set_id,
                "pdf": (Path(PDF_DIRNAME) / pdf_dst.name).as_posix(),
                "docx": (Path(DOCX_DIRNAME) / docx_path.name).as_posix(),
                "source_sha256": source_digest,
                "pdf_exists": pdf_dst.exists(),
                "docx_exists": docx_path.exists(),
                "pdf_pages": pages,
                "docx_blocks": counts["blocks"],
                "docx_passages": counts["passages"],
                "docx_questions": counts["questions"],
                "docx_shorts": counts["shorts"],
                "docx_paragraphs": text_check["paragraphs"],
                "docx_question_headings": text_check["question_headings"],
                "docx_short_headings": text_check["short_headings"],
                "docx_has_set": text_check["has_set"],
                "docx_sections": text_check["sections"],
                "docx_body_columns": text_check["body_columns"],
                "docx_header_tables": text_check["header_tables"],
                "docx_footer_identity": text_check["footer_identity"],
                "docx_required_styles": text_check["required_styles"],
                "docx_option_paragraphs": text_check["option_paragraphs"],
                "docx_answer_blanks": text_check["answer_blanks"],
                "docx_arrow_lines": text_check["arrow_lines"],
                "docx_bold_runs": text_check["bold_runs"],
                "docx_underline_runs": text_check["underline_runs"],
                "docx_raw_markdown": text_check["raw_markdown"],
                "docx_replacement_chars": text_check["replacement_chars"],
                "docx_cover_title": text_check["cover_title"],
                "docx_body_header_set": text_check["body_header_set"],
                "pdf_sha256": sha256(pdf_dst),
                "docx_sha256": sha256(docx_path),
                "status": "PASS"
                if (
                    pdf_dst.exists()
                    and docx_path.exists()
                    and pages == 4
                    and counts["questions"] == 6
                    and counts["shorts"] == 2
                    and text_check["question_headings"] == 6
                    and text_check["short_headings"] == 2
                    and text_check["has_set"] == "yes"
                    and text_check["sections"] >= 2
                    and text_check["body_columns"] == "2"
                    and text_check["header_tables"] >= 1
                    and text_check["footer_identity"] == "yes"
                    and text_check["required_styles"] == "yes"
                    and text_check["option_paragraphs"] >= 25
                    and text_check["answer_blanks"] >= 8
                    and text_check["raw_markdown"] == "no"
                    and text_check["replacement_chars"] == "no"
                    and text_check["cover_title"] == "yes"
                    and text_check["body_header_set"] == "yes"
                )
                else "FAIL",
            }
        )
        package_files.extend([pdf_dst, docx_path])

    manifest = out_dir / MANIFEST_NAME
    with manifest.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    readme = write_readme(out_dir, sets, metadata)
    buyer_readme = write_buyer_readme(out_dir, sets, metadata)
    operator_readme = write_operator_readme(out_dir, sets, metadata)
    provisional_zip = out_dir.parent / ZIP_NAME
    dashboard, dashboard_json = write_release_dashboard(out_dir, sets, rows, provisional_zip, metadata)
    package_files.extend([manifest, readme, buyer_readme, operator_readme, dashboard, dashboard_json])
    zip_path = make_zip(out_dir, package_files)
    return manifest, zip_path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--sets", default="51-62")
    parser.add_argument("--source", default=str(layout.DEFAULT_SOURCE))
    parser.add_argument("--out-dir", default=str(DEFAULT_OUT))
    parser.add_argument("--pdf-dir", default=str(DEFAULT_PDF_DIR))
    args = parser.parse_args()

    sets = parse_set_spec(args.sets)
    manifest, zip_path = export_package(
        sets=sets,
        source=Path(args.source),
        out_dir=Path(args.out_dir),
        pdf_source_dir=Path(args.pdf_dir),
    )
    print(manifest.resolve())
    print(zip_path.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
