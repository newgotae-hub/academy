# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import csv
import json
import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

import export_g1_editable_package as editable
import render_g1_exam_layout as layout


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PACKAGE_DIR = ROOT / "output" / "editable_package"
DEFAULT_REPORT = DEFAULT_PACKAGE_DIR / "g1_mid_editable_package_contract.md"
DEFAULT_JSON = DEFAULT_PACKAGE_DIR / "g1_mid_editable_package_contract.json"
DEFAULT_RECEIPT = DEFAULT_PACKAGE_DIR / editable.RELEASE_RECEIPT_NAME
DEFAULT_RECEIPT_JSON = DEFAULT_PACKAGE_DIR / editable.RELEASE_RECEIPT_JSON_NAME

EXPECTED_COLUMNS = [
    "package_version",
    "release_id",
    "generated_at",
    "release_status",
    "sellable",
    "school",
    "grade",
    "term",
    "exam_scope",
    "set_range_group",
    "docx_edit_policy",
    "pdf_distribution_policy",
    "edit_safe_scope",
    "edit_risky_scope",
    "known_caveat",
    "set",
    "pdf",
    "docx",
    "source_sha256",
    "pdf_exists",
    "docx_exists",
    "pdf_pages",
    "docx_blocks",
    "docx_passages",
    "docx_questions",
    "docx_shorts",
    "docx_paragraphs",
    "docx_question_headings",
    "docx_short_headings",
    "docx_has_set",
    "docx_sections",
    "docx_body_columns",
    "docx_header_tables",
    "docx_footer_identity",
    "docx_required_styles",
    "docx_option_paragraphs",
    "docx_answer_blanks",
    "docx_arrow_lines",
    "docx_bold_runs",
    "docx_underline_runs",
    "docx_raw_markdown",
    "docx_replacement_chars",
    "docx_cover_title",
    "docx_body_header_set",
    "pdf_sha256",
    "docx_sha256",
    "status",
]


def add_check(checks: list[dict[str, str]], scope: str, name: str, ok: bool, detail: str = "") -> None:
    checks.append(
        {
            "scope": scope,
            "name": name,
            "status": "PASS" if ok else "FAIL",
            "detail": detail,
        }
    )


def read_manifest(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        return list(reader.fieldnames or []), list(reader)


def is_inside(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def docx_text(doc: Document) -> str:
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(cell.text for table in section.header.tables for row in table.rows for cell in row.cells)
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def section_column_count(section) -> str:
    cols = section._sectPr.xpath("./w:cols")
    if not cols:
        return ""
    return cols[0].get(qn("w:num"), "")


def inspect_docx(path: Path, set_id: str) -> dict[str, int | str]:
    doc = Document(path)
    text = docx_text(doc)
    style_names = {p.style.name for p in doc.paragraphs}
    return {
        "sections": len(doc.sections),
        "body_columns": section_column_count(doc.sections[1]) if len(doc.sections) > 1 else "",
        "body_header_tables": len(doc.sections[1].header.tables) if len(doc.sections) > 1 else 0,
        "body_footer_identity": "yes"
        if len(doc.sections) > 1 and "광영여자고등학교" in "\n".join(p.text for p in doc.sections[1].footer.paragraphs)
        else "no",
        "required_styles": "yes"
        if {"ExamQuestion", "ExamPassageTitle", "ExamOption", "ExamShort", "ExamSmall"}.issubset(style_names)
        else "no",
        "questions": len(re.findall(r"(?m)^(?:6|7|8|9|10|11)\.\s+\[[0-9.]+점\]", text)),
        "shorts": len(re.findall(r"(?m)^단답형\s+[12]\.\s+\[[0-9.]+점\]", text)),
        "option_paragraphs": sum(1 for p in doc.paragraphs if p.style.name == "ExamOption"),
        "answer_blanks": text.count("________________"),
        "arrows": text.count("->"),
        "raw_markdown": "yes" if "**" in text or "`" in text else "no",
        "replacement_chars": "yes" if "\ufffd" in text else "no",
        "cover_title": "yes" if "2025학년도 1학년" in text and "공통영어1" in text else "no",
        "set_marker": "yes" if f"Set {set_id}" in text or f"Set\n{set_id}" in text else "no",
        "paragraphs": len(doc.paragraphs),
    }


def inspect_docx_archive(path: Path) -> dict[str, int | str]:
    banned_entry_patterns = (
        "word/vbaProject.bin",
        "word/comments.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/embeddings/",
    )
    tracked_tag_re = re.compile(r"<w:(ins|del|moveFrom|moveTo|commentRangeStart|altChunk)(\s|>)")
    external_rels = 0
    banned_entries: list[str] = []
    tracked_count = 0

    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        banned_entries = [name for name in names if any(name.startswith(pattern) for pattern in banned_entry_patterns)]
        for name in names:
            if not name.endswith(".xml") and not name.endswith(".rels"):
                continue
            data = zf.read(name).decode("utf-8", errors="ignore")
            if name.endswith(".rels") and 'TargetMode="External"' in data:
                external_rels += data.count('TargetMode="External"')
            if name.startswith("word/document") or name.startswith("word/header") or name.startswith("word/footer"):
                tracked_count += len(tracked_tag_re.findall(data))

    return {
        "banned_entries": len(banned_entries),
        "banned_entry_sample": ", ".join(banned_entries[:3]),
        "tracked_tags": tracked_count,
        "external_rels": external_rels,
    }


def normalize_for_coverage(text: str) -> str:
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"_+", "____", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def source_coverage(blocks: list[layout.Block], docx_path: Path) -> dict[str, int | str]:
    doc = Document(docx_path)
    full_text = normalize_for_coverage(docx_text(doc))
    expected: list[str] = []
    for block in blocks:
        expected.append(block.title)
        for raw in block.lines:
            line = raw.strip()
            if not line or line == "---":
                continue
            expected.append(line)

    missing: list[str] = []
    for item in expected:
        normalized = normalize_for_coverage(item)
        if normalized and normalized not in full_text:
            missing.append(normalized[:80])

    return {
        "expected_lines": len(expected),
        "missing_lines": len(missing),
        "missing_sample": " | ".join(missing[:3]),
    }


def expected_zip_entries(package_dir: Path, sets: list[str]) -> set[str]:
    entries = {
        editable.MANIFEST_NAME,
        editable.README_NAME,
        editable.BUYER_README_NAME,
        editable.OPERATOR_README_NAME,
        editable.RELEASE_DASHBOARD_NAME,
        editable.RELEASE_DASHBOARD_JSON_NAME,
    }
    for set_id in sets:
        entries.add(f"{editable.PDF_DIRNAME}/g1_mid_set{set_id}_layout.pdf")
        entries.add(f"{editable.DOCX_DIRNAME}/g1_mid_set{set_id}_editable.docx")
    return entries


def audit_package(
    sets: list[str],
    source: Path,
    package_dir: Path,
    expected_pages: int,
) -> tuple[str, list[dict[str, str]], dict[str, object]]:
    checks: list[dict[str, str]] = []
    manifest_path = package_dir / editable.MANIFEST_NAME
    zip_path = package_dir.parent / editable.ZIP_NAME

    add_check(checks, "package", "package_dir_exists", package_dir.exists(), str(package_dir))
    add_check(checks, "package", "manifest_exists", manifest_path.exists(), str(manifest_path))
    add_check(checks, "package", "zip_exists", zip_path.exists(), str(zip_path))
    readme_path = package_dir / editable.README_NAME
    buyer_readme_path = package_dir / editable.BUYER_README_NAME
    operator_readme_path = package_dir / editable.OPERATOR_README_NAME
    dashboard_path = package_dir / editable.RELEASE_DASHBOARD_NAME
    dashboard_json_path = package_dir / editable.RELEASE_DASHBOARD_JSON_NAME
    add_check(checks, "package", "readme_exists", readme_path.exists(), str(readme_path))
    add_check(checks, "package", "buyer_readme_exists", buyer_readme_path.exists(), str(buyer_readme_path))
    add_check(checks, "package", "operator_readme_exists", operator_readme_path.exists(), str(operator_readme_path))
    add_check(checks, "package", "release_dashboard_exists", dashboard_path.exists(), str(dashboard_path))
    add_check(checks, "package", "release_dashboard_json_exists", dashboard_json_path.exists(), str(dashboard_json_path))

    if not manifest_path.exists():
        return "FAIL", checks, {"sets": sets, "package_dir": str(package_dir), "zip": str(zip_path)}

    fieldnames, rows = read_manifest(manifest_path)
    row_by_set = {row.get("set", ""): row for row in rows}
    add_check(checks, "manifest", "columns_exact", fieldnames == EXPECTED_COLUMNS, f"columns={len(fieldnames)}")
    add_check(checks, "manifest", "row_count", len(rows) == len(sets), f"rows={len(rows)} expected={len(sets)}")
    add_check(checks, "manifest", "set_order", [row.get("set") for row in rows] == sets, ",".join(row.get("set", "") for row in rows))
    add_check(checks, "manifest", "all_status_pass", all(row.get("status") == "PASS" for row in rows), "")
    add_check(checks, "manifest", "package_version", all(row.get("package_version") == editable.PACKAGE_VERSION for row in rows), "")
    release_ids = {row.get("release_id", "") for row in rows}
    generated_ats = {row.get("generated_at", "") for row in rows}
    add_check(checks, "manifest", "release_id_single", len(release_ids) == 1 and "" not in release_ids, ",".join(sorted(release_ids)))
    add_check(checks, "manifest", "release_id_format", all(re.match(r"^KY-G1MID-BODY-[0-9]+-[0-9]+-[0-9a-f]{12}$", row.get("release_id", "")) for row in rows), "")
    add_check(checks, "manifest", "generated_at_single", len(generated_ats) == 1 and "" not in generated_ats, ",".join(sorted(generated_ats)))
    add_check(checks, "manifest", "release_status_ship", all(row.get("release_status") == "SHIP" for row in rows), "")
    add_check(checks, "manifest", "sellable_yes", all(row.get("sellable") == "yes" for row in rows), "")
    add_check(checks, "manifest", "policy_columns_populated", all(row.get("docx_edit_policy") and row.get("known_caveat") for row in rows), "")
    add_check(
        checks,
        "manifest",
        "source_hash_consistent",
        all(row.get("source_sha256") == editable.sha256(source) for row in rows),
        "",
    )

    expected_entries = expected_zip_entries(package_dir, sets)
    zip_entry_names: list[str] = []
    zip_integrity = "MISSING"
    if zip_path.exists():
        with zipfile.ZipFile(zip_path) as zf:
            names = set(zf.namelist())
            bad = zf.testzip()
        zip_entry_names = sorted(names)
        zip_integrity = bad or "PASS"
        add_check(checks, "zip", "zip_integrity", bad is None, bad or "PASS")
        add_check(checks, "zip", "zip_entries_exact", names == expected_entries, f"actual={len(names)} expected={len(expected_entries)}")

    if readme_path.exists():
        text = readme_path.read_text(encoding="utf-8")
        add_check(checks, "docs", "readme_release_verdict", "Release Verdict" in text and "SHIP" in text, "")
        add_check(checks, "docs", "readme_points_to_buyer_operator", editable.BUYER_README_NAME in text and editable.OPERATOR_README_NAME in text, "")
    if buyer_readme_path.exists():
        text = buyer_readme_path.read_text(encoding="utf-8")
        add_check(checks, "docs", "buyer_usage_sequence", "수정 후 확인 순서" in text and "PDF로 다시 저장" in text, "")
        add_check(checks, "docs", "buyer_edit_scope", "수정 가능 범위" in text and "주의" in text, "")
    if operator_readme_path.exists():
        text = operator_readme_path.read_text(encoding="utf-8")
        add_check(checks, "docs", "operator_regen_commands", "run_g1_layout_full_verification.py" in text and "audit_g1_editable_package.py" in text, "")
        add_check(checks, "docs", "operator_hold_conditions", "릴리스 보류 조건" in text and "source line coverage" in text, "")
    if dashboard_path.exists() and dashboard_json_path.exists():
        dashboard = dashboard_path.read_text(encoding="utf-8")
        dashboard_json = json.loads(dashboard_json_path.read_text(encoding="utf-8"))
        add_check(checks, "dashboard", "dashboard_status_ship", dashboard_json.get("status") == "SHIP" and "Release Verdict" in dashboard, "")
        add_check(checks, "dashboard", "dashboard_counts_match_manifest", dashboard_json.get("manifest_rows") == len(rows) and dashboard_json.get("pdf_count") == len(rows) and dashboard_json.get("docx_count") == len(rows), "")
        add_check(checks, "dashboard", "dashboard_release_id_matches", len(release_ids) == 1 and dashboard_json.get("release_id") in release_ids and str(dashboard_json.get("release_id", "")) in dashboard, "")
        add_check(checks, "dashboard", "dashboard_sections", all(token in dashboard for token in ["Sellable Scope", "Editable Safety", "Distribution Readiness", "Operator Action"]), "")

    for set_id in sets:
        row = row_by_set.get(set_id, {})
        pdf = package_dir / editable.PDF_DIRNAME / f"g1_mid_set{set_id}_layout.pdf"
        docx = package_dir / editable.DOCX_DIRNAME / f"g1_mid_set{set_id}_editable.docx"
        source_blocks = layout.parse_blocks(layout.extract_set(source, set_id))
        source_questions = sum(1 for block in source_blocks if block.kind == "question")
        source_shorts = sum(1 for block in source_blocks if block.kind == "short")

        add_check(checks, f"set{set_id}", "manifest_row_exists", bool(row), "")
        add_check(checks, f"set{set_id}", "pdf_exists", pdf.exists(), str(pdf))
        add_check(checks, f"set{set_id}", "docx_exists", docx.exists(), str(docx))
        if row:
            pdf_from_manifest = package_dir / row["pdf"]
            docx_from_manifest = package_dir / row["docx"]
            add_check(
                checks,
                f"set{set_id}",
                "manifest_paths_relative",
                not Path(row["pdf"]).is_absolute() and not Path(row["docx"]).is_absolute(),
                f"pdf={row['pdf']} docx={row['docx']}",
            )
            add_check(
                checks,
                f"set{set_id}",
                "manifest_paths_inside_package",
                is_inside(pdf_from_manifest, package_dir) and is_inside(docx_from_manifest, package_dir),
                "",
            )
            add_check(
                checks,
                f"set{set_id}",
                "manifest_paths_exist",
                pdf_from_manifest.exists() and docx_from_manifest.exists(),
                "",
            )
            add_check(checks, f"set{set_id}", "manifest_sha_matches_pdf", pdf.exists() and row.get("pdf_sha256") == editable.sha256(pdf), "")
            add_check(checks, f"set{set_id}", "manifest_sha_matches_docx", docx.exists() and row.get("docx_sha256") == editable.sha256(docx), "")
            add_check(checks, f"set{set_id}", "manifest_counts_match_source", row.get("docx_questions") == str(source_questions) and row.get("docx_shorts") == str(source_shorts), "")

        if pdf.exists():
            pages = len(PdfReader(str(pdf)).pages)
            add_check(checks, f"set{set_id}", "pdf_page_count", pages == expected_pages, f"pages={pages}")

        if docx.exists():
            docx_info = inspect_docx(docx, set_id)
            archive_info = inspect_docx_archive(docx)
            coverage = source_coverage(source_blocks, docx)
            add_check(checks, f"set{set_id}", "docx_two_sections", int(docx_info["sections"]) >= 2, str(docx_info["sections"]))
            add_check(checks, f"set{set_id}", "docx_two_column_body", docx_info["body_columns"] == "2", str(docx_info["body_columns"]))
            add_check(checks, f"set{set_id}", "docx_header_table", int(docx_info["body_header_tables"]) >= 1, str(docx_info["body_header_tables"]))
            add_check(checks, f"set{set_id}", "docx_footer_identity", docx_info["body_footer_identity"] == "yes", "")
            add_check(checks, f"set{set_id}", "docx_required_styles", docx_info["required_styles"] == "yes", "")
            add_check(checks, f"set{set_id}", "docx_question_headings", docx_info["questions"] == source_questions, f"{docx_info['questions']} expected={source_questions}")
            add_check(checks, f"set{set_id}", "docx_short_headings", docx_info["shorts"] == source_shorts, f"{docx_info['shorts']} expected={source_shorts}")
            add_check(checks, f"set{set_id}", "docx_option_paragraphs", int(docx_info["option_paragraphs"]) >= 25, str(docx_info["option_paragraphs"]))
            add_check(checks, f"set{set_id}", "docx_answer_space", int(docx_info["answer_blanks"]) >= 8 and int(docx_info["arrows"]) >= 4, f"blanks={docx_info['answer_blanks']} arrows={docx_info['arrows']}")
            add_check(checks, f"set{set_id}", "docx_no_raw_markdown", docx_info["raw_markdown"] == "no", "")
            add_check(checks, f"set{set_id}", "docx_no_replacement_chars", docx_info["replacement_chars"] == "no", "")
            add_check(checks, f"set{set_id}", "docx_cover_title", docx_info["cover_title"] == "yes", "")
            add_check(checks, f"set{set_id}", "docx_set_marker", docx_info["set_marker"] == "yes", "")
            add_check(checks, f"set{set_id}", "docx_no_banned_entries", int(archive_info["banned_entries"]) == 0, str(archive_info["banned_entry_sample"]))
            add_check(checks, f"set{set_id}", "docx_no_tracked_changes", int(archive_info["tracked_tags"]) == 0, str(archive_info["tracked_tags"]))
            add_check(checks, f"set{set_id}", "docx_no_external_relationships", int(archive_info["external_rels"]) == 0, str(archive_info["external_rels"]))
            add_check(checks, f"set{set_id}", "docx_source_line_coverage", int(coverage["missing_lines"]) == 0, f"expected={coverage['expected_lines']} missing={coverage['missing_lines']} {coverage['missing_sample']}")

    status = "PASS" if all(check["status"] == "PASS" for check in checks) else "FAIL"
    release_id = next(iter(release_ids)) if len(release_ids) == 1 else ""
    generated_at = next(iter(generated_ats)) if len(generated_ats) == 1 else ""
    zip_sha256 = editable.sha256(zip_path) if zip_path.exists() else ""
    payload: dict[str, object] = {
        "contract_version": "g1-editable-package-contract-v2",
        "package_version": editable.PACKAGE_VERSION,
        "release_id": release_id,
        "package_generated_at": generated_at,
        "validated_at": datetime.now().isoformat(timespec="seconds"),
        "source_sha256": editable.sha256(source),
        "manifest_sha256": editable.sha256(manifest_path) if manifest_path.exists() else "",
        "zip_sha256": zip_sha256,
        "zip_integrity": zip_integrity,
        "zip_entry_count": len(zip_entry_names),
        "zip_entry_names": zip_entry_names,
        "status": status,
        "sets": sets,
        "manifest_rows": len(rows),
        "zip_entries_expected": len(expected_entries),
        "expected_pages": expected_pages,
        "package_dir": str(package_dir.resolve()),
        "manifest": str(manifest_path.resolve()),
        "zip": str(zip_path.resolve()),
        "checks": checks,
        "summary": {
            "total": len(checks),
            "pass": sum(1 for check in checks if check["status"] == "PASS"),
            "fail": sum(1 for check in checks if check["status"] == "FAIL"),
        },
    }
    return status, checks, payload


def write_report(path: Path, payload: dict[str, object]) -> None:
    checks = payload["checks"]
    assert isinstance(checks, list)
    summary = payload["summary"]
    assert isinstance(summary, dict)
    lines = [
        "# 광영여고 고1 본문동형 PDF+DOCX 편집패키지 계약검수",
        "",
        f"- status: {payload['status']}",
        f"- release_id: {payload.get('release_id', '')}",
        f"- package_version: {payload.get('package_version', '')}",
        f"- package_generated_at: {payload.get('package_generated_at', '')}",
        f"- validated_at: {payload.get('validated_at', '')}",
        f"- zip_sha256: {payload.get('zip_sha256', '')}",
        f"- sets: {', '.join(payload['sets'])}",
        f"- checks: {summary['pass']} pass / {summary['fail']} fail / {summary['total']} total",
        "",
        "| scope | check | status | detail |",
        "|---|---:|---:|---|",
    ]
    for check in checks:
        lines.append(f"| {check['scope']} | {check['name']} | {check['status']} | {check['detail']} |")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def rebuild_package_zip(package_dir: Path, sets: list[str]) -> Path:
    zip_path = package_dir.parent / editable.ZIP_NAME
    package_files = [package_dir / entry for entry in sorted(expected_zip_entries(package_dir, sets))]
    return editable.make_zip(package_dir, package_files)


def write_dashboard_from_payload(package_dir: Path, payload: dict[str, object]) -> tuple[Path, Path]:
    dashboard_path = package_dir / editable.RELEASE_DASHBOARD_NAME
    dashboard_json_path = package_dir / editable.RELEASE_DASHBOARD_JSON_NAME
    summary = payload["summary"]
    assert isinstance(summary, dict)
    dashboard_payload = {
        "dashboard_version": "g1-editable-release-dashboard-v2",
        "source": "audit_payload",
        "package_version": payload.get("package_version", ""),
        "release_id": payload.get("release_id", ""),
        "package_generated_at": payload.get("package_generated_at", ""),
        "validated_at": "see release receipt",
        "zip_sha256": "see release receipt",
        "manifest_sha256": payload.get("manifest_sha256", ""),
        "zip_integrity": payload.get("zip_integrity", ""),
        "zip_entry_count": payload.get("zip_entry_count", 0),
        "status": "SHIP" if payload.get("status") == "PASS" else "HOLD",
        "sellable": payload.get("status") == "PASS",
        "sets": payload.get("sets", []),
        "manifest_rows": payload.get("manifest_rows", 0),
        "pdf_count": payload.get("manifest_rows", 0),
        "docx_count": payload.get("manifest_rows", 0),
        "zip_entries_expected": payload.get("zip_entries_expected", 0),
        "contract_summary": summary,
    }
    dashboard_json_path.write_text(json.dumps(dashboard_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    dashboard_path.write_text(
        "# 광영여고 고1 본문동형 PDF+DOCX Release Dashboard\n\n"
        "## Release Verdict\n\n"
        f"- status: `{dashboard_payload['status']}`\n"
        f"- sellable: `{'yes' if dashboard_payload['sellable'] else 'no'}`\n"
        f"- package version: `{dashboard_payload['package_version']}`\n"
        f"- release id: `{dashboard_payload['release_id']}`\n"
        f"- source: `audit_payload`\n\n"
        "## Provenance\n\n"
        f"- package generated at: `{dashboard_payload['package_generated_at']}`\n"
        "- validated at: `see release receipt`\n"
        "- zip sha256: `see release receipt`\n\n"
        "## Sellable Scope\n\n"
        f"- PDF: `{dashboard_payload['pdf_count']}` files\n"
        f"- DOCX: `{dashboard_payload['docx_count']}` files\n"
        f"- manifest rows: `{dashboard_payload['manifest_rows']}`\n"
        f"- manifest sha256: `{dashboard_payload['manifest_sha256']}`\n"
        f"- ZIP integrity: `{dashboard_payload['zip_integrity']}`\n"
        f"- ZIP entry count: `{dashboard_payload['zip_entry_count']}`\n"
        f"- expected ZIP entries: `{dashboard_payload['zip_entries_expected']}`\n\n"
        "## Editable Safety\n\n"
        "- DOCX는 Word 또는 한글에서 수정 가능한 편집본입니다.\n"
        "- 수정 후에는 PDF로 다시 저장하고 2단 흐름, 단답형 답란, 마지막 페이지 여백을 확인해야 합니다.\n"
        "- 문항 추가/삭제, 페이지 수 변경, 답란 구조 변경은 재검수 대상입니다.\n\n"
        "## Distribution Readiness\n\n"
        f"- contract checks: `{summary['pass']}` pass / `{summary['fail']}` fail / `{summary['total']}` total\n"
        f"- contract status: `{payload.get('status')}`\n\n"
        "## Operator Action\n\n"
        "- 그대로 판매/배포할 때는 `pdf/`를 사용합니다.\n"
        "- 커스터마이징 판매를 할 때는 `docx/` 수정 후 PDF 재저장본을 따로 검수합니다.\n"
        "- 이 dashboard가 `SHIP`이 아니면 ZIP 판매를 보류합니다.\n",
        encoding="utf-8",
    )
    return dashboard_path, dashboard_json_path


def write_release_receipt(path: Path, json_path: Path, payload: dict[str, object]) -> None:
    receipt = {
        "receipt_version": "g1-editable-release-receipt-v1",
        "release_id": payload.get("release_id", ""),
        "package_version": payload.get("package_version", ""),
        "package_generated_at": payload.get("package_generated_at", ""),
        "validated_at": payload.get("validated_at", ""),
        "status": "SHIP" if payload.get("status") == "PASS" else "HOLD",
        "sellable": payload.get("status") == "PASS",
        "sets": payload.get("sets", []),
        "manifest_rows": payload.get("manifest_rows", 0),
        "zip": payload.get("zip", ""),
        "zip_sha256": payload.get("zip_sha256", ""),
        "manifest_sha256": payload.get("manifest_sha256", ""),
        "zip_integrity": payload.get("zip_integrity", ""),
        "zip_entry_count": payload.get("zip_entry_count", 0),
        "source_sha256": payload.get("source_sha256", ""),
        "contract_version": payload.get("contract_version", ""),
        "contract_summary": payload.get("summary", {}),
    }
    json_path.write_text(json.dumps(receipt, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    path.write_text(
        "# 광영여고 고1 본문동형 PDF+DOCX Release Receipt\n\n"
        f"- release id: `{receipt['release_id']}`\n"
        f"- status: `{receipt['status']}`\n"
        f"- sellable: `{'yes' if receipt['sellable'] else 'no'}`\n"
        f"- package version: `{receipt['package_version']}`\n"
        f"- package generated at: `{receipt['package_generated_at']}`\n"
        f"- validated at: `{receipt['validated_at']}`\n"
        f"- zip sha256: `{receipt['zip_sha256']}`\n"
        f"- manifest sha256: `{receipt['manifest_sha256']}`\n"
        f"- ZIP integrity: `{receipt['zip_integrity']}`\n"
        f"- ZIP entry count: `{receipt['zip_entry_count']}`\n"
        f"- source sha256: `{receipt['source_sha256']}`\n"
        f"- contract: `{receipt['contract_summary']}`\n\n"
        "이 receipt는 ZIP 생성 이후에 작성되는 외부 검수 영수증입니다. ZIP 내부 파일은 자기 자신의 SHA256을 담을 수 없으므로, 최종 ZIP 동일성 확인은 이 receipt의 `zip_sha256` 값을 기준으로 합니다.\n",
        encoding="utf-8",
    )


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--sets", default="51-62")
    parser.add_argument("--source", default=str(layout.DEFAULT_SOURCE))
    parser.add_argument("--package-dir", default=str(DEFAULT_PACKAGE_DIR))
    parser.add_argument("--expected-pages", type=int, default=4)
    parser.add_argument("--report", default=str(DEFAULT_REPORT))
    parser.add_argument("--json", default=str(DEFAULT_JSON))
    parser.add_argument("--receipt", default=str(DEFAULT_RECEIPT))
    parser.add_argument("--receipt-json", default=str(DEFAULT_RECEIPT_JSON))
    args = parser.parse_args()

    sets = editable.parse_set_spec(args.sets)
    initial_status, _initial_checks, initial_payload = audit_package(
        sets=sets,
        source=Path(args.source),
        package_dir=Path(args.package_dir),
        expected_pages=args.expected_pages,
    )
    write_dashboard_from_payload(Path(args.package_dir), initial_payload)
    rebuild_package_zip(Path(args.package_dir), sets)
    status, _checks, payload = audit_package(
        sets=sets,
        source=Path(args.source),
        package_dir=Path(args.package_dir),
        expected_pages=args.expected_pages,
    )

    json_path = Path(args.json)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    write_report(Path(args.report), payload)
    write_release_receipt(Path(args.receipt), Path(args.receipt_json), payload)
    print(json_path.resolve())
    print(Path(args.report).resolve())
    print(Path(args.receipt_json).resolve())
    print(Path(args.receipt).resolve())
    print(f"EDITABLE_PACKAGE_CONTRACT={status}")
    return 0 if status == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
